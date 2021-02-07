import pyttsx3
import datetime
import speech_recognition as sr
import wikipedia
import webbrowser
import os
import random
import time
import smtplib
import docx
from autocorrect import Speller
from pyaudio import PyAudio
import pocketsphinx
import pygame
from pygame.locals import * 


def speak(audio):
    '''
    Speaks the audio(which is text), which is passed in the argument-Ex speak("Hello baby")
    '''
    engine = pyttsx3.init('sapi5')# install this instead #pip install -Iv pyttsx3==2.6 -U
    voices = engine.getProperty('voices')
    # print(voices[1].id)
    rate = engine.getProperty('rate')
    engine.setProperty('rate',198)
    engine.setProperty('voice',voices[1].id)
    engine.say(audio)
    engine.runAndWait()
    # engine.endLoop()
    engine.stop()

# def demo_for_asyncio_loop(audio):
    
#     loop = asyncio.new_event_loop()
    

def speakAndPrint(audio):
    '''
    Speaks and prints the audio(which is text), which is passed in the argumeny-Ex speak("Hello baby")
    '''
    engine = pyttsx3.init('sapi5')# install this instead #pip install -Iv pyttsx3==2.6 -U
    voices = engine.getProperty('voices')
    # print(voices[1].id)
    rate = engine.getProperty('rate')
    engine.setProperty('rate',198)
    engine.setProperty('voice',voices[1].id)
    engine.say(audio)
    engine.runAndWait()
    # engine.endLoop()
    engine.stop()
    print(audio)
    

def wishMe():
    '''
    This functions tells about the current time(whether morning, evening or Afternoon)
    '''
    hour = int(datetime.datetime.now().hour)
    if hour>=0 and hour<12:
        speak("Good Morning!")
    
    elif hour>=12 and hour<=18:
        speak("Good Afternoon!")
    
    else:
        speak("Good Evening!")
        
    
    speak("Hello, my name's Shasha. Please tell me how may I help you")


def takeCommand():
    '''
    Takes microphone input from the user and returns string output.
    ''' 
    r = sr.Recognizer()
    while 1:       
        with sr.Microphone() as source:
            print("Listening...")
            dummy = screen_display(color)
            font = pygame.font.Font('freesansbold.ttf', 32)
            text = font.render("Listening...", True, color, blue)
            screen.blit(text, (200,200))
            pygame.display.update()

            r.pause_threshold = 1 # ctrl + toggle on the fn pause_threshold
            r.phrase_threshold = 0.8
            r.energy_threshold=150
            # engine.runAndWait()# abhi abhi dala
            audio = r.listen(source) 

        try:
            dummy = screen_display(color)
            font = pygame.font.Font('freesansbold.ttf', 32)
            text = font.render("Recognizing...", True, color, blue)
            screen.blit(text, (200,200))
            pygame.display.update()

            print('Recognizing...')
            query = r.recognize_google(audio, language = 'en-IN')
            # query = r.recognize_sphinx(audio, language = 'en-US',grammar=T)
            print(f"User said: {query}\n")
            break

        except Exception as e:
            speakAndPrint('Please repeat Luv')     
            # query = takeCommand()  
        
    
    return query

def sendEmail(to, content):
    server = smtplib.SMTP('smtp.gmail.com', 587)
    server.ehlo()
    server.starttls()
    with open("pass.txt", "r") as f:
        global email,password1
        email = f.readline()
        password1 = f.readline()
    # server.login(f'{email}, 'luv')
    server.login(user = email,password = password1)
    print("1")
    server.sendmail('luvkaranwal123@gmail.com', to, content)
    print("2")
    server.close()


def sleep2(query):
    test = query.split()
    for i in range(0,len(test)):
        try:
            test[i] = int(test[i])
            time.sleep(test[i])

        except Exception as e:
            pass     


# For Screen | Textbox | Text
def screen_display(color):
    screen.fill(bg_color)
    font = pygame.font.Font('freesansbold.ttf', 32)
    if color == color_active:
        condition='Active'
    else:        
        condition='Inactive'        
    text = font.render(condition, True, color, green)
    coordinates = (600,300)
    screen.blit(text, coordinates)
    input_box = text.get_rect().move(coordinates)

    return input_box

if __name__ == "__main__":
    # speak("I'm in love with luv")

    
    # chrome_path="C:\\Users\\Love Karnval\\AppData\\Local\\Google\\Chrome\\Application\\chrome.exe"
    # webbrowser.register('chrome', None,webbrowser.BackgroundBrowser(chrome_path))
    # controller = webbrowser.get('chrome')


    #----------------------pygame defaults ------------------------------
    
    pygame.init()
    screen = pygame.display.set_mode((0,0), flags = pygame.FULLSCREEN)
    bg_color = pygame.Color(255, 129, 170)
    white = (255, 255, 255)
    green = (0, 255, 0)
    blue = (0, 0, 128)
    color_active = pygame.Color('dodgerblue2')
    color_inactive = pygame.Color('lightskyblue3')
    pygame.display.set_caption('Shasha - The Voice Assistant') 
    capture=True
    active = False
    clock = pygame.time.Clock()
    color=color_inactive

    input_box = screen_display(color)
    pygame.display.update()

    #-------------------------------------------------------------------

    wishMe()
    a=0
    while capture:
        for event in pygame.event.get():
            if event.type == pygame.QUIT:
                capture = False
            elif event.type == KEYDOWN and event.key == K_q:
                capture = False
            elif event.type == pygame.MOUSEBUTTONDOWN:
                if input_box.collidepoint(event.pos):
                    active = True
                    color = color_active
                    input_box = screen_display(color)
                    pygame.display.update()
                    print('Active...')
                else:
                    active = False
                    print('Not Active...')
            
            if active:
                if a>0:
                    speak("So, what would you like me to do now, luv?") 
                else:
                    a+=1
                
                screen.fill(bg_color)

                query = takeCommand().lower()

                # Logic for executing tasks based on query.
                if 'sleep' in query:
                    sleep2(query) 

                
                elif 'wikipedia' in query:
                    speak('Searching Wikipedia...')
                    query = query.replace("wikipedia","")
                    results = wikipedia.summary(query, sentences = 2)
                    speak("According to wikipedia")
                    print(results)
                    speak(results)
                

                elif 'open youtube' in query:
                    print("Here in youtube")
                    urL='https://www.youtube.com'
                    # controller.open(urL)
                    webbrowser.open(urL)

                elif 'open google' in query:
                    url = 'https://www.google.com'
                    controller.open(url)

                elif 'tujhe pata hai' in query:
                    url = 'https://www.youtube.com/makejokeof'
                    # controller.open(url)
                    webbrowser.open(urL)   
                
                elif 'stackoverflow' in query:
                    url = 'https://www.stackoverflow.com'
                    # controller.open(url)
                    webbrowser.open(urL)

                elif 'search' and 'youtube' in query:
                    query = query.replace(' in youtube',"")
                    query = query.replace('search ',"")
                    query = query.replace(" "," + ") 
                    print(query)
                    url = "http://www.youtube.com/results?search_query =" + query
                    print(url)
                    # driver.get("http://www.youtube.com/results?search_query =" + '+'.join(subquery)) 
                    # controller.open(url)
                    webbrowser.open(urL) 
                

                elif 'open music'in query:
                    speak("What do you want to do")
                    subquery = takeCommand().lower()
                    test = subquery.split()  
                    # print(test)  
                    music_dir = 'C:\\Users\\Love Karnval\\Desktop\\SONGS'     
                    songs = os.listdir(music_dir)
                    # print(list(songs))
                    a=0

                    while True: 
                        if(a>0): 
                            query = takeCommand().lower()    
                            test = query.split()

                        if 'play' or 'change' or 'next' in test:
                            i= random.randint(0,30)
                            # print(i)
                            os.startfile(os.path.join(music_dir, songs[i]))
                            a+=1    

                        elif 'close music' in query:
                            os.system("taskkill /IM \"Music.UI.exe\" /F")   
                            
                            break

                        # elif 'stop' in test:
                        #     print("Printed immediately.")
                        #     time.sleep(90)
                        #     print("Printed after 90 seconds.")
                        #     a+=1
                        #     break
                elif 'the time' in query:
                    strTime = datetime.datetime.now().strftime("%H:%M:%S")
                    speak(f"Sir, the time is {strTime}")   

                elif 'open attack' in query:
                    codePath = "G:\\dxd 2\\shingeki no kyoujin 1 sub"     
                    os.startfile(codePath)
                    
                elif 'send email' in query:
                    try:
                        speak("What should I send ?")
                        content = takeCommand()
                        to = "luvkaranwal123@gmail.com"
                        sendEmail(to, content)
                        speak("Email has been sent!")
                        print("Email has been sent!")            
                    
                    except Exception as e:
                        print(e)
                        speak("Sorry luv, couldn't send the email")    
                
                elif 'open word' in query:
                    speakAndPrint("What operation do you want to perform")
                    print('1. open a doc')
                    print('2. create a doc')
                    check_command = takeCommand().lower()
                    
                    if 'create' in check_command:
                        '''
                        create a doc file
                        '''
                        doc = docx.Document()
                    
                        # while True:
                        #     if temp == '':
                        #         speakAndPrint('Sorry, but couldn\'t recognize the paragraph.')
                        #         continue
                        #     break    
                        speakAndPrint("Please speak the paragraph.")
                        temp = takeCommand() # Takes a paragraph as input
                        
                        check = Speller(lang = 'en')
                        para = check(temp)
                        print("After auto-Correction:\n",para)
                        a = 0
                        while True:
                            speakAndPrint("Do you want to keep the initial paragraph or the one after auto-correction?")
                            print("1.autocorrected-Yes")
                            print("2.original-No")
                            speak("Please choose from the above choices as you require, and in the mean time I would be asleep for some time to let you decide, which one to choose.")
                            time.sleep(30)
                            speakAndPrint('So what have you decided?')
                            check_command = takeCommand().lower() 
                            
                            if 'yes' in check_command:
                                break    # The edited one.
                                    
                            elif 'no' in check_command:
                                
                                para = temp # The unedited one            
                                # Since para is passed in the .add_paragraph
                                break

                            else:
                                speak('Please speak a correct response')
                                continue

                        while True:
                            speakAndPrint("Do you want to replace a statement or word ?")
                            print("1. Yes")
                            print("2. No")
                            time.sleep(2)
                            check_command = takeCommand().lower()
                            if  'yes' in check_command:
                                speakAndPrint("Please speak what you want to replace in the paragraph")
                                replacing_stat = takeCommand()
                                speakAndPrint("Searching...")
                                if (replacing_stat in para.lower()) or (replacing_stat in para.lower()) or replacing_stat in para:
                                    try:
                                        speakAndPrint("match found...")
                                        speakAndPrint("Please speak what you want to put in the place of '"+ replacing_stat +" ' ")
                                        replaced_stat= takeCommand()
                                        speakAndPrint("Trying to replace...")
                                        para = para.replace(replacing_stat, replaced_stat)
                                        speakAndPrint("replaced succesfully")
                                        speakAndPrint("The new paragraph is:")
                                        print(para)
                                        speak("Please review you paragraph and in the mean time I would be asleep for some time.")
                                        time.sleep(20)
                                    except Exception as e:
                                        speak("Sorry.. but an unexpected error occurred. Let me ask you once again.")
                                        continue    
                                else:
                                    speakAndPrint("Sorry, but it seems the term you said isn't present in the paragraph or has a gra1matical misatake")
                                    continue
                            elif 'no' in check_command:
                                break
                            else:
                                speak('Please speak a correct response.')    

                        doc_para = doc.add_paragraph(para)
                        while True:
                            speak('Please tell me what would you like to name your document ?')
                            
                            doc_name = takeCommand()
                            speakAndPrint('Do you want to name your file as '+ doc_name)
                            speak('1.Yes')
                            speak('1.No')
                            time.sleep(2)
                            choice = takeCommand().lower()
                            if 'yes' in choice:
                                doc.save('E:\\khud se kiya python apun\\luvpy\\Jarvis\\documents in jarvis\\' + doc_name + ".docx")
                                speakAndPrint('Saved succesfully')
                                break
                            elif 'no' in choice:
                                speak('Please speak')
                                continue
                            else:
                                speak('Please speak up a correct response.')
                                continue
                    
                    elif 'open' in check_command:
                        '''
                        open a doc from a given list.
                        '''
                        while True:
                            doc_name=''
                            try:
                        
                                speakAndPrint("Please speak a document from the below list that you would like to open.")
                                arr = os.listdir('documents in jarvis')    
                                i=1
                                for items in arr:
                                    print(i,".",items)
                                    i+=1
                                speakAndPrint('So...,which file would you like to open')
                                time.sleep(3)
                                doc_name = takeCommand().lower()
                                doc = os.startfile('E:\\khud se kiya python apun\\luvpy\\Jarvis\\documents in jarvis\\'+ doc_name +".docx")
                                speakAndPrint('Would you like to open another file?')
                                print('1.Yes')
                                print('2.No')
                                choice = takeCommand().lower()
                                if 'yes' in choice:
                                    continue
                                elif 'no' in choice:
                                    break
                                else:
                                    print('Sorry but, couldn\'t identify the input given or maybe its an incorrect response.\nPlease let me ask you once again.')     
                        
                            except Exception as e:
                                speakAndPrint('It seems as '+ doc_name +'.docx doesn\'t exists or cannot be opened due to some reasons. So, please try speaking the name again.')
                        
                        # f = open('naya2.docx', 'rb')

                    elif 'close' in check_command:
                        '''
                        Close the document
                        '''
                        os.system("taskkill /IM \"WINWORD.exe\" /F") 

                elif 'sasha quit' in query.lower():
                    speak('It was nice speaking up with you luv, c u later.')
                    break
                active = False
                if active == False:
                    color = color_inactive
                    input_box = screen_display(color)
                    pygame.display.update()

        clock.tick(60)
        pygame.display.update()